# -*- coding: utf-8 -*-
from __future__ import annotations

import os
import re
import json
import time
import math
import shutil
import logging
from dataclasses import dataclass
from pathlib import Path
from datetime import datetime
from collections import defaultdict
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import Any, Dict, List, Optional, Tuple

import requests
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
from tqdm import tqdm
from colorama import Fore, Style, init as colorama_init

from ..config.loader import load_merged_config
from ..common.logging_utils import setup_logger

colorama_init(autoreset=True)

PROJECT = "Humanities Doc Toolkit - Sorter"
VERSION = "0.1"


# ----------------------------
# JSON helpers (兼容模型返回 ```json ... ```) [2]
# ----------------------------
def _clean_json_text(text: str) -> str:
    t = (text or "").strip()
    if t.startswith("```"):
        t = re.sub(r"^```[a-zA-Z]*\n", "", t).strip()
        t = t.rstrip("`").strip()
    return t

def _extract_json_obj(text: str) -> Optional[dict]:
    t = _clean_json_text(text)
    if not t:
        return None
    if not (t.startswith("{") and t.endswith("}")):
        m = re.search(r"\{.*\}", t, re.DOTALL)
        if m:
            t = m.group(0)
    try:
        obj = json.loads(t)
        return obj if isinstance(obj, dict) else None
    except Exception:
        return None


# ----------------------------
# API statistics (来自 v5.1 结构) [2]
# ----------------------------
class APIStatistics:
    def __init__(self):
        self.stats = {
            "total_calls": 0,
            "successful_calls": 0,
            "failed_calls": 0,
            "tokens_used": 0,
            "api_key_usage": defaultdict(int),
            "service_usage": defaultdict(int),
            "batch_stats": {"first_round_batches": 0, "second_round_calls": 0},
        }

    def record(self, api_key_name: str, service: str, success: bool, tokens: int = 0, call_type: str = "unknown"):
        self.stats["total_calls"] += 1
        self.stats["tokens_used"] += int(tokens or 0)
        self.stats["api_key_usage"][api_key_name] += 1
        self.stats["service_usage"][service] += 1
        if success:
            self.stats["successful_calls"] += 1
        else:
            self.stats["failed_calls"] += 1

        if call_type == "batch_first_round":
            self.stats["batch_stats"]["first_round_batches"] += 1
        elif call_type == "single_content":
            self.stats["batch_stats"]["second_round_calls"] += 1

    def summary(self) -> Dict[str, Any]:
        out = dict(self.stats)
        out["api_key_usage"] = dict(out["api_key_usage"])
        out["service_usage"] = dict(out["service_usage"])
        out["batch_stats"] = dict(out["batch_stats"])
        return out


# ----------------------------
# AI Service selector (v5.1) [2]
# ----------------------------
class AIServiceSelector:
    def __init__(self, merged_cfg: Dict[str, Any]):
        self.cfg = merged_cfg

    def _services(self) -> Dict[str, Any]:
        services = self.cfg.get("ai_services", {}).get("services", {})
        # 兼容旧配置中 google=gemini [1]
        if "gemini" not in services and "google" in services:
            services["gemini"] = services["google"]
        # 兼容旧配置中 moonshot=kimi [1]
        if "kimi" not in services and "moonshot" in services:
            services["kimi"] = services["moonshot"]
        return services

    def available_services(self) -> List[Dict[str, Any]]:
        out = []
        for name, scfg in self._services().items():
            if not isinstance(scfg, dict) or not scfg.get("enabled", False):
                continue
            keys = [k for k in scfg.get("api_keys", []) if isinstance(k, dict) and k.get("enabled", False) and str(k.get("key", "")).strip()]
            if not keys:
                continue
            out.append({"name": name, "model": scfg.get("model", ""), "api_count": len(keys)})
        return out

    def show_service_menu(self) -> Optional[str]:
        av = self.available_services()
        print(f"{Fore.CYAN}🤖 请选择AI服务提供商{Style.RESET_ALL}")
        print("=" * 50)

        if not av:
            print(f"{Fore.RED}❌ 未检测到可用的AI服务配置{Style.RESET_ALL}")
            print("请检查 global.yaml 中 ai_services.services.* 的 enabled 与 api_keys")
            return None

        for i, s in enumerate(av, 1):
            print(f"{Fore.WHITE}{i}. {s['name']}{Style.RESET_ALL}")
            print(f"   模型: {s['model']}")
            print(f"   API密钥数量: {s['api_count']} 个\n")

        while True:
            choice = input(f"{Fore.YELLOW}请选择 (1-{len(av)}): {Style.RESET_ALL}").strip()
            if choice.isdigit():
                idx = int(choice) - 1
                if 0 <= idx < len(av):
                    selected = av[idx]["name"]
                    print(f"{Fore.GREEN}✓ 已选择: {selected}{Style.RESET_ALL}")
                    return selected
            print(f"{Fore.RED}❌ 无效选择,请输入 1-{len(av)}{Style.RESET_ALL}")


# ----------------------------
# Filtering mode selector (从配置读取，完全对齐 v5.1 filtering_modes) [1][2]
# ----------------------------
class FilteringModeSelector:
    def __init__(self, tool_cfg: Dict[str, Any]):
        self.cfg = tool_cfg
        self.modes = self._load_modes()

    def _load_modes(self) -> Dict[str, Any]:
        modes = self.cfg.get("document_sorting", {}).get("filtering_modes", {})
        if not isinstance(modes, dict) or not modes:
            raise ValueError("sorter.yaml 缺少 document_sorting.filtering_modes [1]")
        return modes

    def show_mode_menu(self) -> Dict[str, Any]:
        print(f"{Fore.CYAN}🎯 请选择筛选精度模式{Style.RESET_ALL}")
        print("=" * 60)

        keys = list(self.modes.keys())
        for i, k in enumerate(keys, 1):
            m = self.modes[k]
            name = m.get("name", k)
            desc = m.get("description", "")
            fr = m.get("first_round_threshold", 0.3)
            sr = m.get("second_round_threshold", 0.7)
            bs = m.get("batch_size", 8)
            ue = m.get("enable_universal_enhancement", True)
            rs = m.get("run_second_round", True)
            print(f"{i}. {name}")
            if desc:
                print(f"   {desc}")
            print(f"   第一轮阈值: {fr}")
            print(f"   第二轮阈值: {sr}")
            print(f"   批次大小: {bs}")
            print(f"   普遍化增强: {'开启' if ue else '关闭'}")
            print(f"   第二轮: {'执行' if rs else '跳过(仅第一轮)'}\n")

        default_mode = self.cfg.get("ai", {}).get("default_mode", "balanced")
        default_idx = keys.index(default_mode) + 1 if default_mode in keys else 1

        while True:
            c = input(f"{Fore.YELLOW}请选择 (1-{len(keys)}) [默认{default_idx}]: {Style.RESET_ALL}").strip()
            if not c:
                c = str(default_idx)
            if c.isdigit():
                idx = int(c) - 1
                if 0 <= idx < len(keys):
                    k = keys[idx]
                    m = dict(self.modes[k])
                    m["key"] = k
                    print(f"{Fore.GREEN}✓ 已选择: {m.get('name', k)}{Style.RESET_ALL}")
                    return m
            print(f"{Fore.RED}❌ 无效选择,请输入 1-{len(keys)}{Style.RESET_ALL}")


# ----------------------------
# Document analyzer (对齐 v5.1: 智能类型分析 + 多格式提取) [1][2]
# ----------------------------
class EnhancedDocumentAnalyzer:
    def __init__(self, cfg: Dict[str, Any], logger: logging.Logger):
        self.cfg = cfg
        self.logger = logger
        da = cfg.get("document_analysis", {})
        self.paper_max = int(da.get("page_thresholds", {}).get("paper_max", 50))
        self.book_min = int(da.get("page_thresholds", {}).get("book_min", 100))
        ex = da.get("extraction", {})
        self.paper_pages = int(ex.get("paper_pages", 8))
        self.book_pages = int(ex.get("book_pages", 15))
        self.book_toc_pages = int(ex.get("book_toc_pages", 10))

    def analyze_document_type(self, path: Path) -> Dict[str, Any]:
        page_count = self._get_page_count(path)
        if page_count <= self.paper_max:
            return {"page_count": page_count, "document_type": "paper",
                    "extract_strategy": {"type": "paper", "pages_to_extract": self.paper_pages, "token_budget": 2500, "analyze_toc": False}}
        if page_count >= self.book_min:
            return {"page_count": page_count, "document_type": "book",
                    "extract_strategy": {"type": "book", "pages_to_extract": self.book_pages, "toc_pages": self.book_toc_pages, "token_budget": 3500, "analyze_toc": True}}
        return {"page_count": page_count, "document_type": "medium",
                "extract_strategy": {"type": "medium", "pages_to_extract": min(12, max(1, page_count // 2)), "token_budget": 3000, "analyze_toc": False}}

    def _get_page_count(self, path: Path) -> int:
        suf = path.suffix.lower()
        try:
            if suf == ".pdf":
                import PyPDF2
                with open(path, "rb") as f:
                    r = PyPDF2.PdfReader(f)
                    return len(r.pages)
            if suf in [".docx", ".doc"]:
                from docx import Document
                doc = Document(path)
                para = len([p for p in doc.paragraphs if p.text.strip()])
                return max(1, para // 15)
            # text-like
            txt = path.read_text(encoding="utf-8", errors="ignore")
            return max(1, len(txt) // 3000)
        except Exception:
            return 1

    def extract_smart_content(self, path: Path, strategy: Dict[str, Any]) -> str:
        suf = path.suffix.lower()
        if suf == ".pdf":
            return self._extract_pdf(path, strategy)
        if suf in [".docx", ".doc"]:
            return self._extract_docx(path, strategy)
        if suf in [".txt", ".md", ".rtf"]:
            return self._extract_text(path, strategy)
        return ""

    def _extract_pdf(self, path: Path, strategy: Dict[str, Any]) -> str:
        max_chars = int(strategy.get("token_budget", 2500))
        parts: List[str] = []
        try:
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                total = len(pdf.pages)
                if strategy.get("type") == "book" and strategy.get("analyze_toc"):
                    toc_pages = int(strategy.get("toc_pages", 10))
                    toc = []
                    for i in range(min(toc_pages, total)):
                        t = pdf.pages[i].extract_text() or ""
                        if t.strip():
                            toc.append(t)
                        if sum(len(x) for x in toc) > max_chars // 2:
                            break
                    parts.append("===目录与前言===\n" + "\n".join(toc)[: max_chars // 2])

                    mid_start = min(toc_pages + 5, max(0, total - 5))
                    sample = []
                    for i in range(mid_start, min(mid_start + 3, total)):
                        t = pdf.pages[i].extract_text() or ""
                        if t.strip():
                            sample.append(t)
                        if sum(len(x) for x in sample) > max_chars // 2:
                            break
                    if sample:
                        parts.append("===内容样本===\n" + "\n".join(sample)[: max_chars // 2])
                else:
                    n = int(strategy.get("pages_to_extract", 8))
                    for i in range(min(n, total)):
                        t = pdf.pages[i].extract_text() or ""
                        if t.strip():
                            parts.append(t)
                        if sum(len(x) for x in parts) > max_chars:
                            break
        except Exception:
            # 回退 PyPDF2（会在某些PDF上产生 unknown widths 噪声，需屏蔽） 
            try:
                import PyPDF2
                import logging as _logging
                from contextlib import redirect_stderr
                from io import StringIO

                # 1) 降低 PyPDF2 logger（有些版本走 logging）
                _logging.getLogger("PyPDF2").setLevel(_logging.ERROR)

                # 2) 屏蔽写入 stderr 的噪声（unknown widths 往往从这里出来）
                _buf = StringIO()
                with redirect_stderr(_buf):
                    with open(path, "rb") as f:
                        r = PyPDF2.PdfReader(f)
                        n = int(strategy.get("pages_to_extract", 8))
                        for i in range(min(n, len(r.pages))):
                            t = r.pages[i].extract_text() or ""
                            if t.strip():
                                parts.append(t)
                            if sum(len(x) for x in parts) > max_chars:
                                break
            except Exception:
                parts = ["无法提取PDF内容"]
        return ("\n".join(parts))[:max_chars]

    def _extract_docx(self, path: Path, strategy: Dict[str, Any]) -> str:
        max_chars = int(strategy.get("token_budget", 2000))
        try:
            from docx import Document
            doc = Document(path)
            parts = []
            limit = int(strategy.get("pages_to_extract", 8)) * 15
            for i, p in enumerate(doc.paragraphs):
                if i >= limit:
                    break
                t = (p.text or "").strip()
                if t:
                    parts.append(t)
                if sum(len(x) for x in parts) > max_chars:
                    break
            return ("===DOCX===\n" + "\n".join(parts))[:max_chars]
        except Exception:
            return "无法提取Word文档内容"

    def _extract_text(self, path: Path, strategy: Dict[str, Any]) -> str:
        max_chars = int(strategy.get("token_budget", 2000))
        try:
            import chardet
            raw = path.read_bytes()
            enc = (chardet.detect(raw).get("encoding") or "utf-8")
            txt = path.read_text(encoding=enc, errors="ignore")
            return ("===TEXT===\n" + txt)[:max_chars]
        except Exception:
            return "无法提取文本内容"


# ----------------------------
# API Manager (完整实现 v5.1 的三个调用入口) [2]
# ----------------------------
class EnhancedAPIManager:
    def __init__(self, merged_cfg: Dict[str, Any], selected_service: str, stats: APIStatistics, logger: logging.Logger):
        self.cfg = merged_cfg
        self.stats = stats
        self.logger = logger
        self.service = selected_service

        services = self.cfg.get("ai_services", {}).get("services", {})
        if "gemini" not in services and "google" in services:
            services["gemini"] = services["google"]
        if "kimi" not in services and "moonshot" in services:
            services["kimi"] = services["moonshot"]

        self.service_cfg = services.get(selected_service)
        if not self.service_cfg:
            raise ValueError(f"未找到服务配置: {selected_service}")

        self.keys = [k for k in self.service_cfg.get("api_keys", []) if k.get("enabled", False) and str(k.get("key", "")).strip()]
        if not self.keys:
            raise ValueError(f"{selected_service} 没有可用密钥")

        self.idx = 0
        self.session = requests.Session()
        retry = Retry(total=2, backoff_factor=0.5, status_forcelist=[429, 500, 502, 503, 504])
        adapter = HTTPAdapter(max_retries=retry, pool_connections=10, pool_maxsize=10)
        self.session.mount("https://", adapter)
        self.session.mount("http://", adapter)

        self.logger.info(f"API管理器就绪: {selected_service}, keys={len(self.keys)}")

    def _next_key(self) -> Dict[str, Any]:
        k = self.keys[self.idx % len(self.keys)]
        self.idx += 1
        return k

    def _estimate_tokens(self, prompt: str, answer: str) -> int:
        return len(prompt) // 4 + len(answer) // 4

    def _call(self, prompt: str, key_info: Dict[str, Any], call_type: str) -> str:
        service_name = self.service
        base_url = str(self.service_cfg.get("base_url", "")).rstrip("/")
        model = self.service_cfg.get("model")
        timeout = int(self.service_cfg.get("timeout", 60))
        max_tokens = int(self.service_cfg.get("max_tokens", 3000))
        temperature = float(self.service_cfg.get("temperature", 0.1))

        # Claude [1]
        if service_name == "claude" or "anthropic" in base_url:
            headers = {"x-api-key": key_info["key"], "anthropic-version": "2023-06-01", "content-type": "application/json"}
            payload = {"model": model, "max_tokens": max_tokens, "temperature": temperature, "messages": [{"role": "user", "content": prompt}]}
            t0 = time.time()
            try:
                r = self.session.post(base_url, headers=headers, json=payload, timeout=timeout)
                r.raise_for_status()
                data = r.json()
                text = ""
                if isinstance(data.get("content"), list) and data["content"]:
                    text = data["content"][0].get("text", "")
                tokens = self._estimate_tokens(prompt, text)
                self.stats.record(key_info.get("name", "key"), self.service, True, tokens, call_type=call_type)
                return text
            except Exception as e:
                self.stats.record(key_info.get("name", "key"), self.service, False, 0, call_type=call_type)
                raise e
            finally:
                _ = time.time() - t0

        # Gemini（旧配置里是 google）[1]
        if service_name == "gemini" or "generativelanguage.googleapis.com" in base_url:
            url = base_url
            if "key=" not in url:
                url = url + ("&" if "?" in url else "?") + f"key={key_info['key']}"
            payload = {"contents": [{"parts": [{"text": prompt}]}],
                       "generationConfig": {"temperature": temperature, "maxOutputTokens": max_tokens}}
            t0 = time.time()
            try:
                r = self.session.post(url, json=payload, timeout=timeout)
                r.raise_for_status()
                data = r.json()
                text = ""
                try:
                    text = data["candidates"][0]["content"]["parts"][0]["text"]
                except Exception:
                    text = ""
                tokens = self._estimate_tokens(prompt, text)
                self.stats.record(key_info.get("name", "key"), self.service, True, tokens, call_type=call_type)
                return text
            except Exception as e:
                self.stats.record(key_info.get("name", "key"), self.service, False, 0, call_type=call_type)
                raise e
            finally:
                _ = time.time() - t0

        # OpenAI-compatible（deepseek/openai/kimi/glm/qwen 等）[1][2]
        headers = {"Content-Type": "application/json", "Authorization": f"Bearer {key_info['key']}"}
        payload = {
            "model": model,
            "messages": [{"role": "user", "content": prompt}],
            "max_tokens": max_tokens,
            "temperature": temperature,
        }
        t0 = time.time()
        try:
            r = self.session.post(base_url, headers=headers, json=payload, timeout=timeout)
            r.raise_for_status()
            data = r.json()
            text = data["choices"][0]["message"]["content"]
            tokens = self._estimate_tokens(prompt, text)
            self.stats.record(key_info.get("name", "key"), self.service, True, tokens, call_type=call_type)
            return text
        except Exception as e:
            self.stats.record(key_info.get("name", "key"), self.service, False, 0, call_type=call_type)
            raise e
        finally:
            _ = time.time() - t0

    # —— v5.1: 普遍化增强 —— [2]
    def generate_universal_enhancement_prompt(self, user_query: str) -> str:
        prompt = (
            "你是语义扩展专家。请为以下研究需求生成普遍化概念扩展和搜索提示词增强:\n\n"
            f"用户研究需求:{user_query}\n\n"
            "请生成:\n"
            "1. 核心概念识别\n"
            "2. 更广泛相关概念(如:康德种族思想 → 西方种族观念史)\n"
            "3. 中英文相关词汇\n"
            "4. 学科交叉概念\n\n"
            "返回格式(保持简洁):\n"
            "【普遍化语义扩展】\n核心概念:...\n广泛概念:...\n相关词汇:...\n交叉概念:...\n\n"
            "【搜索增强指导】\n注意:专门主题可能作为更广泛主题的章节出现,请评估这种包含关系。"
        )
        key = self._next_key()
        try:
            text = self._call(prompt, key, call_type="enhancement")
            return f"\n{text}\n"
        except Exception:
            return "\n【使用默认语义扩展】\n请考虑专门主题与广泛主题的包含关系。\n"

    # —— v5.1: 第一轮批量 —— [2]
    def batch_first_round_call(self, file_names: List[str], user_query: str, universal_enhancement: str) -> Dict[str, Any]:
        key = self._next_key()
        file_list = "\n".join([f"{i+1}. {name}" for i, name in enumerate(file_names)])
        prompt = f"""你是专业的学术文献评估专家。请基于文件名评估与研究需求的相关性。

研究需求:{user_query}

{universal_enhancement}

文档列表(共{len(file_names)}个):
{file_list}

评估要求:
1. 考虑直接匹配和普遍化关系(如专门主题可能是更广泛主题的一部分)
2. 评估语义相关性,包括跨语言理解
3. 返回所有文档的评分,格式严格按照JSON

返回格式:
{{"files": [
    {{"name": "完整文件名", "score": 0.85, "reason": "详细分析理由"}}
]}}

评分标准:
- 0.9-1.0: 高度相关
- 0.7-0.8: 相关性强
- 0.5-0.6: 中等相关
- 0.3-0.4: 轻微相关
- 0.0-0.2: 不相关

请确保返回JSON包含所有{len(file_names)}个文档的评估结果。"""
        try:
            text = self._call(prompt, key, call_type="batch_first_round")
            return {"success": True, "content": text, "api_key": key.get("name", "key"), "batch_size": len(file_names)}
        except Exception as e:
            return {"success": False, "error": str(e), "api_key": key.get("name", "key"), "batch_size": len(file_names)}

    # —— v5.1: 第二轮单文档 —— [2]
    def single_content_analysis_call(self, document_info: Dict[str, Any], user_query: str, universal_enhancement: str) -> Dict[str, Any]:
        key = self._next_key()
        prompt = f"""请基于文档实际内容深度分析与研究需求的匹配度。

研究需求:{user_query}

{universal_enhancement}

文档信息:
- 文件名:{document_info['name']}
- 文档类型:{document_info.get('document_type', '未知')}
- 页数:{document_info.get('page_count', '未知')}页

{document_info.get('content_preview', '无内容预览')}

分析要求:
1. 基于实际内容(非仅文件名)进行深度分析
2. 考虑普遍化关系和间接相关性
3. 评估内容对研究需求的实际价值

返回JSON格式:
{{"score": 0.85, "reason": "基于内容的详细匹配分析", "content_highlights": "关键内容要点"}}"""
        try:
            text = self._call(prompt, key, call_type="single_content")
            return {"success": True, "content": text, "api_key": key.get("name", "key")}
        except Exception as e:
            return {"success": False, "error": str(e), "api_key": key.get("name", "key")}


# ----------------------------
# Index scanner (对齐 v5.1: recursive_scan + supported_formats + skip_hidden + size filter) [1][2]
# ----------------------------
class DocumentIndex:
    def __init__(self, cfg: Dict[str, Any], logger: logging.Logger):
        self.cfg = cfg
        self.logger = logger

    def create_index(self, directory: Path) -> List[Dict[str, Any]]:
        scan = self.cfg.get("document_sorting", {}).get("scanning", {})
        recursive = bool(scan.get("recursive_scan", True))
        supported = set([s.lower() for s in scan.get("supported_formats", [".pdf"])])
        skip_hidden = bool(scan.get("skip_hidden", True))
        min_kb = int(scan.get("min_file_size_kb", 10))
        max_mb = int(scan.get("max_file_size_mb", 500))
        max_depth = int(scan.get("max_depth", 10))

        docs: List[Dict[str, Any]] = []

        def _walk(p: Path, depth: int):
            if depth > max_depth:
                return
            try:
                for item in p.iterdir():
                    if skip_hidden and item.name.startswith("."):
                        continue
                    if item.is_dir():
                        if recursive:
                            _walk(item, depth + 1)
                        continue
                    if item.is_file():
                        suf = item.suffix.lower()
                        if suf not in supported:
                            continue
                        st = item.stat()
                        size_kb = st.st_size / 1024
                        size_mb = st.st_size / 1024 / 1024
                        if size_kb < min_kb:
                            continue
                        if max_mb > 0 and size_mb > max_mb:
                            continue
                        docs.append({
                            "path": str(item),
                            "name": item.name,
                            "stem": item.stem,
                            "suffix": suf,
                            "size": st.st_size,
                            "size_mb": round(size_mb, 2),
                        })
            except Exception:
                return

        _walk(directory, 0)
        docs.sort(key=lambda d: d.get("size", 0))
        return docs


# ----------------------------
# Semantic filter (两轮筛选 + fast_first_round) [2][1]
# ----------------------------
class EnhancedSemanticFilter:
    def __init__(self, cfg: Dict[str, Any], api_manager: EnhancedAPIManager, mode: Dict[str, Any], stats: APIStatistics, logger: logging.Logger):
        self.cfg = cfg
        self.api_manager = api_manager
        self.mode = mode
        self.stats = stats
        self.logger = logger

        self.analyzer = EnhancedDocumentAnalyzer(cfg, logger)

        self.first_threshold = float(mode.get("first_round_threshold", 0.3))
        self.second_threshold = float(mode.get("second_round_threshold", 0.7))
        self.batch_size = int(mode.get("batch_size", 8))
        self.enable_universal = bool(mode.get("enable_universal_enhancement", True))
        self.run_second_round = bool(mode.get("run_second_round", True))

        # 线程配置来自 sorter.yaml runtime 或 v5.1 performance.threading.max_workers [1]
        self.max_threads = int(cfg.get("performance", {}).get("threading", {}).get("max_workers", 4))

        print(f"{Fore.CYAN}🎯 语义筛选器就绪 - {mode.get('name', mode.get('key','mode'))}{Style.RESET_ALL}")
        print(f"   第一轮阈值: {self.first_threshold}, 批次: {self.batch_size}/批")
        print(f"   第二轮阈值: {self.second_threshold}, 线程: {self.max_threads}")
        print(f"   普遍化增强: {'开启' if self.enable_universal else '关闭'}")
        print(f"   第二轮: {'执行' if self.run_second_round else '跳过(仅第一轮)'}")

    def first_round_batch_filtering(self, documents: List[Dict[str, Any]], user_query: str) -> List[Dict[str, Any]]:
        print(f"\n{Fore.YELLOW}🎯 第一轮批量文件名筛选({len(documents)} 个文档){Style.RESET_ALL}")
        if not documents:
            return []

        universal = ""
        if self.enable_universal:
            print("   🧠 AI生成普遍化语义增强...")
            universal = self.api_manager.generate_universal_enhancement_prompt(user_query)
        else:
            universal = "\n【语义匹配指导】\n请进行精确语义匹配,避免过度普遍化。\n"

        batches = [documents[i:i+self.batch_size] for i in range(0, len(documents), self.batch_size)]
        print(f"   📦 批次配置:{self.batch_size}个/批,共{len(batches)}批")

        passed: List[Dict[str, Any]] = []
        failed_batches = 0

        with ThreadPoolExecutor(max_workers=min(self.max_threads, len(batches))) as ex:
            futs = {}
            for bi, batch in enumerate(batches, 1):
                names = [d["name"] for d in batch]
                futs[ex.submit(self.api_manager.batch_first_round_call, names, user_query, universal)] = (bi, batch)

            with tqdm(total=len(batches), desc=f"第一轮批量({self.batch_size}/批)", bar_format="{l_bar}{bar}| {n_fmt}/{total_fmt} [{elapsed}<{remaining}]") as bar:
                for fut in as_completed(futs):
                    bi, batch = futs[fut]
                    try:
                        resp = fut.result()
                        if resp.get("success"):
                            obj = _extract_json_obj(resp.get("content", "")) or {}
                            files = obj.get("files", []) if isinstance(obj.get("files", []), list) else []
                            name2doc = {d["name"]: d for d in batch}
                            for item in files:
                                nm = item.get("name")
                                if nm in name2doc:
                                    score = float(item.get("score", 0) or 0)
                                    if score >= self.first_threshold:
                                        d = name2doc[nm]
                                        d["first_round_score"] = score
                                        d["first_round_reason"] = item.get("reason", "")
                                        d["batch_num"] = bi
                                        d["api_key"] = resp.get("api_key", "")
                                        passed.append(d)
                        else:
                            failed_batches += 1
                    except Exception:
                        failed_batches += 1
                    finally:
                        bar.update(1)

        passed.sort(key=lambda x: x.get("first_round_score", 0), reverse=True)
        print(f"{Fore.GREEN}✓ 第一轮完成:{len(passed)}/{len(documents)} 个文档通过{Style.RESET_ALL}")
        if failed_batches:
            print(f"{Fore.YELLOW}⚠ 第一轮失败批次: {failed_batches}{Style.RESET_ALL}")
        return passed

    def second_round_content_filtering(self, documents: List[Dict[str, Any]], user_query: str) -> List[Dict[str, Any]]:
        print(f"\n{Fore.YELLOW}🔬 第二轮内容分析({len(documents)}个文档){Style.RESET_ALL}")
        if not documents:
            return []

        universal = ""
        if self.enable_universal:
            universal = self.api_manager.generate_universal_enhancement_prompt(user_query)
        else:
            universal = "\n【精确内容分析】\n请进行精确内容匹配,避免过度普遍化。\n"

        # 预处理：类型分析 + 内容提取（对应 v5.1 的预处理进度条）[2]
        processed: List[Dict[str, Any]] = []
        for d in tqdm(documents, desc="预处理", leave=True, bar_format="{l_bar}{bar}| {n_fmt}/{total_fmt}"):
            analysis = self.analyzer.analyze_document_type(Path(d["path"]))
            d.update(analysis)
            content = self.analyzer.extract_smart_content(Path(d["path"]), analysis["extract_strategy"])
            d["content_preview"] = content
            processed.append(d)

        passed: List[Dict[str, Any]] = []
        failed = 0

        with ThreadPoolExecutor(max_workers=self.max_threads) as ex:
            futs = {ex.submit(self.api_manager.single_content_analysis_call, d, user_query, universal): d for d in processed}
            with tqdm(total=len(processed), desc="AI分析", bar_format="{l_bar}{bar}| {n_fmt}/{total_fmt}") as bar:
                for fut in as_completed(futs):
                    d = futs[fut]
                    try:
                        resp = fut.result()
                        if resp.get("success"):
                            obj = _extract_json_obj(resp.get("content", "")) or {}
                            score = float(obj.get("score", 0) or 0)
                            if score >= self.second_threshold:
                                d["second_round_score"] = score
                                d["second_round_reason"] = obj.get("reason", "")
                                d["content_highlights"] = obj.get("content_highlights", "")
                                d["api_key_used"] = resp.get("api_key", "")
                                passed.append(d)
                        else:
                            failed += 1
                    except Exception:
                        failed += 1
                    finally:
                        bar.update(1)

        passed.sort(key=lambda x: x.get("second_round_score", 0), reverse=True)
        if failed:
            print(f"{Fore.YELLOW}⚠ 第二轮分析失败: {failed}个文档{Style.RESET_ALL}")
        print(f"{Fore.GREEN}✓ 第二轮完成: {len(passed)}/{len(documents)}个文档通过{Style.RESET_ALL}")
        return passed


# ----------------------------
# Folder naming (对齐 v5.1 folder_naming) [1]
# ----------------------------
class FolderNamer:
    def __init__(self, cfg: Dict[str, Any]):
        fn = cfg.get("folder_naming", {})
        self.auto_generate = bool(fn.get("auto_generate", True))
        self.max_name_length = int(fn.get("max_name_length", 30))
        self.add_timestamp = bool(fn.get("add_timestamp", True))
        self.timestamp_format = fn.get("timestamp_format", "%m%d_%H%M")
        self.forbidden = set(fn.get("forbidden_chars", ['<', '>', ':', '"', '/', '\\', '|', '?', '*']))

    def make(self, user_query: str) -> str:
        base = user_query.strip()
        if not base:
            base = "AI智能分拣"
        # 简单抽取前3个词（与 v5.1 generate_folder_name 相似）[2]
        base = base.replace("的", " ").replace("和", " ")
        words = [w for w in re.split(r"\s+", base) if w]
        base_name = "_".join(words[:3]) if words else "AI智能分拣"
        base_name = "".join(ch for ch in base_name if ch not in self.forbidden)
        base_name = base_name[: self.max_name_length]

        if self.add_timestamp:
            ts = datetime.now().strftime(self.timestamp_format)
            base_name = f"{base_name}_{ts}"
        return base_name[: self.max_name_length]


# ----------------------------
# Report generator (对齐 v5.1: 智能分拣详细报告.txt + API统计 + 两轮评分理由) [2]
# ----------------------------
class ReportWriter:
    def __init__(self, cfg: Dict[str, Any], stats: APIStatistics, logger: logging.Logger):
        self.cfg = cfg
        self.stats = stats
        self.logger = logger

    def write(self, documents: List[Dict[str, Any]], target_folder: Path, user_query: str, mode: Dict[str, Any]):
        log_cfg = self.cfg.get("logging", {})
        if not bool(log_cfg.get("generate_detailed_report", True)):
            return

        report_name = log_cfg.get("report_filename", "智能分拣详细报告.txt")
        report_path = target_folder / report_name
        s = self.stats.summary()

        try:
            with report_path.open("w", encoding="utf-8") as f:
                f.write(f"Smart Document Sorter v5.1 - 核心逻辑移植版（工具链）\n")
                f.write("=" * 60 + "\n\n")
                f.write(f"用户研究需求: {user_query}\n")
                f.write(f"筛选模式: {mode.get('name', mode.get('key','mode'))}\n")
                f.write(f"分拣时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"结果文件夹: {target_folder}\n")
                f.write(f"最终文档数量: {len(documents)}\n")
                f.write(f"分析方法: 两轮筛选（文件名批量 + 内容深度）+ 可选普遍化增强\n\n")

                f.write("API调用统计:\n")
                f.write("-" * 30 + "\n")
                f.write(f"总调用次数: {s['total_calls']}\n")
                f.write(f"成功调用: {s['successful_calls']}\n")
                f.write(f"失败调用: {s['failed_calls']}\n")
                f.write(f"Token使用量(估算): {s['tokens_used']:,}\n")
                f.write(f"第一轮批次: {s['batch_stats']['first_round_batches']}\n")
                f.write(f"第二轮单分析: {s['batch_stats']['second_round_calls']}\n\n")

                f.write("API密钥使用分布:\n")
                for k, c in sorted(s["api_key_usage"].items(), key=lambda x: x[1], reverse=True):
                    f.write(f" - {k}: {c}\n")
                f.write("\n")

                f.write("文档分析结果明细:\n")
                f.write("-" * 60 + "\n\n")
                for i, doc in enumerate(documents, 1):
                    f.write(f"【{i:3d}】 {doc['name']}\n")
                    f.write(f"     文件大小: {doc.get('size_mb', 0):.2f} MB\n")
                    f.write(f"     文档类型: {doc.get('document_type', '未知')}\n")
                    f.write(f"     页数估算: {doc.get('page_count', '未知')} 页\n")

                    if "first_round_score" in doc:
                        f.write(f"     第一轮评分: {doc.get('first_round_score', 0):.2f}\n")
                        f.write(f"     文件名分析: {str(doc.get('first_round_reason',''))[:120]}...\n")

                    if "second_round_score" in doc:
                        f.write(f"     第二轮评分: {doc.get('second_round_score', 0):.2f}\n")
                        f.write(f"     内容分析: {str(doc.get('second_round_reason',''))[:160]}...\n")
                        if doc.get("content_highlights"):
                            f.write(f"     内容要点: {str(doc.get('content_highlights',''))[:160]}...\n")

                    f.write(f"     原始路径: {doc['path']}\n")
                    f.write("\n" + "-" * 50 + "\n\n")

            print(f"{Fore.GREEN}✓ 详细报告已生成: {report_path}{Style.RESET_ALL}")
        except Exception as e:
            self.logger.warning(f"生成报告失败: {e}")


# ----------------------------
# Sorter main (交互 + 两轮 + copy) [2][1]
# ----------------------------
class SorterEngine:
    def __init__(self, global_config: str, tool_config: str):
        self.cfg = load_merged_config(global_config, tool_config)
        self.logger = setup_logger("hdt-sorter", str(self.cfg.get("logging", {}).get("level", "INFO")))
        self.stats = APIStatistics()

    def _paths(self) -> Tuple[str, str]:
        # 输入/输出：遵循 v5.1 “运行时输入 input，输出默认 Desktop”风格 [1][2]
        input_dir = input(f"{Fore.CYAN}📂 源文件夹路径: {Style.RESET_ALL}").strip().strip('"')
        if not input_dir:
            raise ValueError("源文件夹路径不能为空")
        if not Path(input_dir).expanduser().exists():
            raise ValueError("源文件夹路径不存在")

        default_output = str((Path.cwd() / "output").resolve())
        out_dir = input(f"{Fore.CYAN}📁 输出目录 (默认: {default_output}): {Style.RESET_ALL}").strip().strip('"') or default_output
        Path(out_dir).expanduser().mkdir(parents=True, exist_ok=True)
        return str(Path(input_dir).expanduser()), str(Path(out_dir).expanduser())

    def _user_query(self) -> str:
        while True:
            q = input(f"{Fore.CYAN}🔍 研究需求: {Style.RESET_ALL}").strip()
            if q:
                return q
            print(f"{Fore.RED}❌ 需求不能为空{Style.RESET_ALL}")

    def _thread_config(self) -> int:
        # v5.1: ai_services.threading_config.max_threads (运行时可自定义) [1][2]
        th_cfg = self.cfg.get("ai_services", {}).get("threading_config", {})
        default_threads = int(th_cfg.get("max_threads", self.cfg.get("runtime", {}).get("max_threads_default", 4)))

        while True:
            s = input(f"{Fore.CYAN}⚙️  最大线程数 (默认 {default_threads}): {Style.RESET_ALL}").strip()
            if not s:
                return default_threads
            if s.isdigit():
                v = int(s)
                if 1 <= v <= 32:
                    return v
            print(f"{Fore.RED}❌ 线程数必须在1-32之间{Style.RESET_ALL}")

    def _apply_thread_policy(self, selected_service: str, requested_threads: int) -> int:
        # 依据 v5.1 api_pool_config.buffer_ratio 的冗余思想 [1]
        buffer_ratio = float(self.cfg.get("ai_services", {}).get("api_pool_config", {}).get("buffer_ratio", 1.5))
        services = self.cfg.get("ai_services", {}).get("services", {})
        if "gemini" not in services and "google" in services:
            services["gemini"] = services["google"]
        if "kimi" not in services and "moonshot" in services:
            services["kimi"] = services["moonshot"]

        keys = [k for k in services.get(selected_service, {}).get("api_keys", []) if k.get("enabled", False) and str(k.get("key","")).strip()]
        available_keys = len(keys)
        if available_keys <= 0:
            return 1

        required = math.ceil(requested_threads * buffer_ratio)
        if available_keys < required:
            # 自动降级线程：floor(keys / ratio)
            safe_threads = max(1, int(available_keys // buffer_ratio))
            print(f"{Fore.YELLOW}⚠ 密钥数量不足以满足 {buffer_ratio}x 冗余：threads={requested_threads} 需要keys≈{required}，当前keys={available_keys}。将自动降为 {safe_threads} 线程。{Style.RESET_ALL}")
            return safe_threads

        return requested_threads

    def _select_service(self) -> str:
        show = bool(self.cfg.get("ai", {}).get("show_service_selection", True))
        selector = AIServiceSelector(self.cfg)
        if show:
            svc = selector.show_service_menu()
            if not svc:
                raise ValueError("未选择AI服务")
            return svc
        # 默认 active_service [1]
        svc = self.cfg.get("ai_services", {}).get("active_service", "deepseek")
        return svc

    def _select_mode(self) -> Dict[str, Any]:
        show = bool(self.cfg.get("ai", {}).get("show_mode_selection", True))
        selector = FilteringModeSelector(self.cfg)
        if show:
            return selector.show_mode_menu()
        default_mode = self.cfg.get("ai", {}).get("default_mode", "balanced")
        modes = self.cfg.get("document_sorting", {}).get("filtering_modes", {})
        m = dict(modes.get(default_mode, list(modes.values())[0]))
        m["key"] = default_mode
        return m

    def _copy_documents(self, docs: List[Dict[str, Any]], target: Path) -> Dict[str, Any]:
        target.mkdir(parents=True, exist_ok=True)
        ok, fail = [], []

        with tqdm(total=len(docs), desc="复制文件", bar_format="{l_bar}{bar}| {n_fmt}/{total_fmt}") as bar:
            for d in docs:
                try:
                    src = Path(d["path"])
                    dst = target / src.name
                    if dst.exists():
                        stem, suf = dst.stem, dst.suffix
                        c = 1
                        while dst.exists():
                            dst = target / f"{stem}_{c}{suf}"
                            c += 1
                    shutil.copy2(src, dst)
                    ok.append({"source": str(src), "target": str(dst), "doc": d})
                except Exception as e:
                    fail.append({"path": d.get("path"), "error": str(e)})
                finally:
                    bar.update(1)

        return {"success": ok, "failed": fail}

    def run_interactive(self) -> int:
        from ..__about__ import __title__, __version__, __author__, __email__, __github__

        print(f"{Fore.MAGENTA}{'='*60}{Style.RESET_ALL}")
        print(f"{Fore.MAGENTA}🧠 {__title__} - Sorter v{__version__}{Style.RESET_ALL}")
        print(f"{Fore.CYAN}👨‍💻 作者: {__author__} | {__email__}{Style.RESET_ALL}")
        print(f"{Fore.CYAN}🔗 {__github__}{Style.RESET_ALL}")
        print(f"{Fore.WHITE}提示:{Style.RESET_ALL}")
        print(f"{Fore.WHITE}- 默认操作是 copy（保留原文件，降低误操作风险）[6]{Style.RESET_ALL}")
        print(f"{Fore.WHITE}- 两轮筛选：文件名批量 + 内容深度（可选 fast_first_round 仅第一轮）[6]{Style.RESET_ALL}")
        print(f"{Fore.MAGENTA}{'='*60}{Style.RESET_ALL}")
        import logging
        logging.getLogger("PyPDF2").setLevel(logging.ERROR)
        selected_service = self._select_service()
        mode = self._select_mode()

        source_dir, output_dir = self._paths()
        user_query = self._user_query()

        requested_threads = self._thread_config()
        max_threads = self._apply_thread_policy(selected_service, requested_threads)

        # 写入 runtime 性能配置（与 v5.1 “运行时可自定义线程”一致）[1][2]
        self.cfg.setdefault("performance", {}).setdefault("threading", {})["max_workers"] = max_threads

        print(f"\n{Fore.CYAN}📋 分拣任务配置:{Style.RESET_ALL}")
        print(f"   源目录: {Path(source_dir).name}")
        print(f"   需求: {user_query}")
        print(f"   筛选模式: {mode.get('name', mode.get('key','mode'))}")
        print(f"   最大线程: {max_threads}")
        print(f"   输出: {Path(output_dir).name}")
        print(f"   操作: copy\n")

        start = time.time()

        indexer = DocumentIndex(self.cfg, self.logger)
        docs = indexer.create_index(Path(source_dir))
        if not docs:
            print(f"{Fore.RED}❌ 未找到支持的文档{Style.RESET_ALL}")
            return 1

        api_manager = EnhancedAPIManager(self.cfg, selected_service, self.stats, self.logger)
        semantic_filter = EnhancedSemanticFilter(self.cfg, api_manager, mode, self.stats, self.logger)

        # 第一轮
        first_pass = semantic_filter.first_round_batch_filtering(docs, user_query)
        if not first_pass:
            print(f"{Fore.RED}❌ 第一轮筛选无结果{Style.RESET_ALL}")
            return 1

        # 第二轮（可选：fast_first_round 跳过）[1]
        if semantic_filter.run_second_round:
            second_pass = semantic_filter.second_round_content_filtering(first_pass, user_query)
            final_docs = second_pass if second_pass else first_pass[:15]
            if not second_pass:
                print(f"{Fore.YELLOW}⚠ 第二轮无结果，使用第一轮Top结果{Style.RESET_ALL}")
        else:
            final_docs = first_pass[: min(200, len(first_pass))]
            print(f"{Fore.YELLOW}⚡ 已选择仅第一轮模式：将直接使用第一轮结果（Top {len(final_docs)}）{Style.RESET_ALL}")

        # 输出文件夹
        folder_name = FolderNamer(self.cfg).make(user_query)
        target_folder = Path(output_dir) / folder_name

        results = self._copy_documents(final_docs, target_folder)

        # 报告
        ReportWriter(self.cfg, self.stats, self.logger).write(final_docs, target_folder, user_query, mode)

        elapsed = time.time() - start
        self._show_summary(results, target_folder, elapsed, mode)
        return 0

    def _show_summary(self, results: Dict[str, Any], target_folder: Path, elapsed: float, mode: Dict[str, Any]):
        ok = len(results.get("success", []))
        fail = len(results.get("failed", []))
        s = self.stats.summary()

        print(f"\n{Fore.MAGENTA}🎉 智能分拣完成!{Style.RESET_ALL}")
        print("=" * 60)
        print(f"{Fore.CYAN}📊 结果统计:{Style.RESET_ALL}")
        print(f"   成功分拣: {ok} 个文档")
        print(f"   失败: {fail} 个文档")
        print(f"   处理用时: {elapsed:.1f} 秒")
        if elapsed > 0:
            print(f"   平均速度: {ok/elapsed:.2f} 文档/秒")
        print(f"   筛选模式: {mode.get('name', mode.get('key','mode'))}")

        print(f"\n{Fore.YELLOW}🤖 API调用统计:{Style.RESET_ALL}")
        if s["total_calls"] > 0:
            print(f"   总调用次数: {s['total_calls']}")
            print(f"   成功率: {s['successful_calls']/s['total_calls']*100:.1f}%")
        print(f"   Token使用(估算): {s['tokens_used']:,}")
        print(f"   第一轮批次: {s['batch_stats']['first_round_batches']}")
        print(f"   第二轮单分析: {s['batch_stats']['second_round_calls']}")

        if s["api_key_usage"]:
            print(f"\n{Fore.GREEN}🔑 API密钥使用分布:{Style.RESET_ALL}")
            for k, c in sorted(s["api_key_usage"].items(), key=lambda x: x[1], reverse=True):
                print(f"   {k}: {c} 次")

        print(f"\n{Fore.BLUE}📁 保存位置: {target_folder}{Style.RESET_ALL}")
        print(f"{Fore.BLUE}📄 详细报告: {target_folder}/智能分拣详细报告.txt{Style.RESET_ALL}")

        try:
            if input(f"\n{Fore.CYAN}打开结果文件夹? (y/N): {Style.RESET_ALL}").lower() == "y":
                import subprocess, platform
                if platform.system() == "Windows":
                    subprocess.run(f'explorer "{target_folder}"', shell=True)
                elif platform.system() == "Darwin":
                    subprocess.run(["open", str(target_folder)])
                else:
                    subprocess.run(["xdg-open", str(target_folder)])
        except Exception:
            pass